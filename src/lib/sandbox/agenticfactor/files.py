"""
AgenticFactor SDK — File Parsing Module
Parse PDFs, DOCX, CSV, Excel, and text files.
"""

import csv
import io
import json
import os
from typing import List, Dict, Optional


def parse_pdf(file_path_or_data, max_pages: int = 100) -> str:
    """
    Parse a PDF file and extract text content.
    
    Args:
        file_path_or_data: File path string, bytes content, or base64 string
        max_pages: Maximum pages to extract
    
    Returns:
        Extracted text content
    """
    try:
        import PyPDF2
    except ImportError:
        try:
            import subprocess
            subprocess.check_call(["pip", "install", "PyPDF2"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            import PyPDF2
        except Exception:
            return "[ERROR: PyPDF2 not available. Cannot parse PDF.]"
    
    if isinstance(file_path_or_data, str) and os.path.exists(file_path_or_data):
        with open(file_path_or_data, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text_parts = []
            for i, page in enumerate(reader.pages[:max_pages]):
                text_parts.append(page.extract_text() or "")
            return "\n\n".join(text_parts)
    elif isinstance(file_path_or_data, bytes):
        reader = PyPDF2.PdfReader(io.BytesIO(file_path_or_data))
        text_parts = []
        for i, page in enumerate(reader.pages[:max_pages]):
            text_parts.append(page.extract_text() or "")
        return "\n\n".join(text_parts)
    elif isinstance(file_path_or_data, str):
        # Try base64
        import base64
        try:
            data = base64.b64decode(file_path_or_data)
            return parse_pdf(data, max_pages)
        except Exception:
            return file_path_or_data  # Return as-is if not a valid PDF
    
    return "[ERROR: Unsupported input type for PDF parsing]"


def parse_docx(file_path_or_data) -> str:
    """Parse a DOCX file and extract text."""
    try:
        from docx import Document
    except ImportError:
        try:
            import subprocess
            subprocess.check_call(["pip", "install", "python-docx"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            from docx import Document
        except Exception:
            return "[ERROR: python-docx not available]"
    
    if isinstance(file_path_or_data, str) and os.path.exists(file_path_or_data):
        doc = Document(file_path_or_data)
    elif isinstance(file_path_or_data, bytes):
        doc = Document(io.BytesIO(file_path_or_data))
    else:
        return "[ERROR: Unsupported input type]"
    
    paragraphs = [p.text for p in doc.paragraphs]
    
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            paragraphs.append(" | ".join(cells))
    
    return "\n".join(paragraphs)


def parse_csv(file_path_or_data, delimiter: str = ",") -> List[List[str]]:
    """Parse a CSV file and return as 2D list."""
    if isinstance(file_path_or_data, str) and os.path.exists(file_path_or_data):
        with open(file_path_or_data, "r", newline="", encoding="utf-8") as f:
            reader = csv.reader(f, delimiter=delimiter)
            return list(reader)
    elif isinstance(file_path_or_data, (str, bytes)):
        text = file_path_or_data if isinstance(file_path_or_data, str) else file_path_or_data.decode("utf-8")
        reader = csv.reader(io.StringIO(text), delimiter=delimiter)
        return list(reader)
    return []


def parse_excel(file_path_or_data, sheet_name: str = None) -> List[List]:
    """Parse an Excel file and return as 2D list."""
    try:
        import openpyxl
    except ImportError:
        try:
            import subprocess
            subprocess.check_call(["pip", "install", "openpyxl"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            import openpyxl
        except Exception:
            return [["ERROR: openpyxl not available"]]
    
    if isinstance(file_path_or_data, str) and os.path.exists(file_path_or_data):
        wb = openpyxl.load_workbook(file_path_or_data, read_only=True)
    elif isinstance(file_path_or_data, bytes):
        wb = openpyxl.load_workbook(io.BytesIO(file_path_or_data), read_only=True)
    else:
        return [["ERROR: Unsupported input type"]]
    
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
    return [[cell.value for cell in row] for row in ws.iter_rows()]


def read_text(file_path: str, encoding: str = "utf-8") -> str:
    """Read a text file."""
    with open(file_path, "r", encoding=encoding) as f:
        return f.read()
